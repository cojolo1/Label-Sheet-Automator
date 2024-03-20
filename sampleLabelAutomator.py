from docx import Document
from docx.shared import Pt
import pandas as pd
import os
import sys
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class sample:
    def __init__(self, studyName, sampleID, lotNo, construct, volume, processStep, date):
        self.studyName = studyName
        self.sampleID = sampleID
        self.lotNo = lotNo
        self.construct = construct
        self.volume = volume
        self.processStep = processStep
        self.date = date
    def addSample(self, dict, noOfReplicates, startIndex):
        for i in range(noOfReplicates):
            dict["<StudyName" + str(startIndex  + i + 1) + ">"] = self.studyName
            dict["<SampleID" + str(startIndex + i + 1)  + ">"] = self.sampleID
            dict["<LotNo" + str(startIndex + i + 1)  + ">"] = self.lotNo
            dict["<Construct" + str(startIndex + i + 1)  + ">"] = self.construct
            dict["<Volume" + str(startIndex + i + 1)  + ">"] = self.volume + "µL " + str(self.date)
            dict["<ProcessStep" + str(startIndex + i + 1)  + ">"] = self.processStep
        return dict
def loadSampleInfo(excelSheet, volume):
    df = pd.read_excel(excelSheet)
    listOfSamples = []
    for index, row in df.iterrows():
        studyName = row["Study Name"]
        sampleID = row["Sample ID"]
        lotNo = row["Lot No"]
        construct = row["Construct"]
        processStep = row["Process Step"]
        date = row["Date"]
        newSample = sample(studyName, sampleID, lotNo, construct, volume, processStep, date)
        listOfSamples.append(newSample)
    return listOfSamples
def compileDict(listOfSamples, noOfReplicates):
    startingDict = {}
    for i in range(len(listOfSamples)):
        newDict = listOfSamples[i].addSample(startingDict, noOfReplicates, i * noOfReplicates)
        startingDict.update(newDict)
    return startingDict
def get_font_info(run):
    font = run.font
    return {
        "name": font.name,
        "size": font.size,
    }
def set_font_info(run, font_info):
    font = run.font
    font.name = font_info["name"]
    font.size = font_info["size"]

def fill_template(template_path, output_path, replacements):
    # Load the template document
    doc = Document(template_path)
    # Replace placeholders in the document
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                font_info = get_font_info(paragraph.runs[0])
                paragraph.text = paragraph.text.replace(key, value)
                set_font_info(paragraph.runs[0], font_info)
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        font_info = get_font_info(cell.paragraphs[0].runs[0])
                        cell.text = cell.text.replace(key, str(value))
                        set_font_info(cell.paragraphs[0].runs[0], font_info)
    # Save the filled-out document
    doc.save(output_path)
# Example usage
if __name__ == "__main__":
    script_path = sys.argv[0]
    folder_path = os.path.dirname(script_path)

    excel_filename = os.path.join(folder_path, "testExcel.xlsx")
    word_filename = os.path.join(folder_path, "labelSheetTemplate.docx")
    output_filename = os.path.join(folder_path, "TFF_CV_Study Pt 2_2.docx")

    listOfSamples = loadSampleInfo(excel_filename, "500")
    replacements = compileDict(listOfSamples, 2)
    fill_template(word_filename, output_filename, replacements)

